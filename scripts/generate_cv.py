#!/usr/bin/env python3
"""Generate polished EN/ZH academic CVs for Ziyi Sun from homepage + prior resumes."""

from __future__ import annotations

from pathlib import Path

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_LINE_SPACING, WD_TAB_ALIGNMENT
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Cm, Pt, RGBColor, Twips

OUT_DIR = Path(__file__).resolve().parents[1] / "resumes"
DATE_TAG = "20260808"


def set_run_font(run, *, size=10.5, bold=False, italic=False, east_asia="SimSun", ascii_font="Times New Roman"):
    run.bold = bold
    run.italic = italic
    run.font.size = Pt(size)
    run.font.color.rgb = RGBColor(0x1A, 0x1A, 0x1A)
    run.font.name = ascii_font
    rpr = run._element.get_or_add_rPr()
    rfonts = rpr.get_or_add_rFonts()
    rfonts.set(qn("w:ascii"), ascii_font)
    rfonts.set(qn("w:hAnsi"), ascii_font)
    rfonts.set(qn("w:cs"), ascii_font)
    rfonts.set(qn("w:eastAsia"), east_asia)


def set_paragraph_format(p, *, before=0, after=0, line=1.05, space_after_pt=None):
    pf = p.paragraph_format
    pf.space_before = Pt(before)
    pf.space_after = Pt(after if space_after_pt is None else space_after_pt)
    pf.line_spacing_rule = WD_LINE_SPACING.MULTIPLE
    pf.line_spacing = line


def add_bottom_border(paragraph):
    pPr = paragraph._p.get_or_add_pPr()
    pBdr = OxmlElement("w:pBdr")
    bottom = OxmlElement("w:bottom")
    bottom.set(qn("w:val"), "single")
    bottom.set(qn("w:sz"), "12")
    bottom.set(qn("w:space"), "1")
    bottom.set(qn("w:color"), "2F2F2F")
    pBdr.append(bottom)
    pPr.append(pBdr)


def configure_page(doc):
    section = doc.sections[0]
    section.page_width = Cm(21.0)
    section.page_height = Cm(29.7)
    section.top_margin = Cm(1.2)
    section.bottom_margin = Cm(1.2)
    section.left_margin = Cm(1.6)
    section.right_margin = Cm(1.6)


def add_heading_line(doc, text, east_asia="SimHei"):
    p = doc.add_paragraph()
    set_paragraph_format(p, before=8, after=2, line=1.0)
    run = p.add_run(text)
    set_run_font(run, size=11.5, bold=True, east_asia=east_asia)
    add_bottom_border(p)
    return p


def add_lr_line(doc, left, right, *, bold_left=True, size=10.5, east_asia="SimSun"):
    p = doc.add_paragraph()
    set_paragraph_format(p, before=2, after=0, line=1.05)
    tab_stops = p.paragraph_format.tab_stops
    tab_stops.add_tab_stop(Cm(17.8), WD_TAB_ALIGNMENT.RIGHT)
    run_l = p.add_run(left)
    set_run_font(run_l, size=size, bold=bold_left, east_asia=east_asia)
    run_tab = p.add_run("\t")
    set_run_font(run_tab, size=size, east_asia=east_asia)
    run_r = p.add_run(right)
    set_run_font(run_r, size=size, bold=False, east_asia=east_asia)
    return p


def add_body(doc, text, *, size=10, italic=False, bold=False, before=0, after=1, east_asia="SimSun"):
    p = doc.add_paragraph()
    set_paragraph_format(p, before=before, after=after, line=1.05)
    run = p.add_run(text)
    set_run_font(run, size=size, italic=italic, bold=bold, east_asia=east_asia)
    return p


def add_bullet(doc, text, *, size=10, east_asia="SimSun"):
    p = doc.add_paragraph()
    set_paragraph_format(p, before=0, after=1, line=1.05)
    p.paragraph_format.left_indent = Cm(0.35)
    p.paragraph_format.first_line_indent = Cm(-0.35)
    run = p.add_run("• " + text)
    set_run_font(run, size=size, east_asia=east_asia)
    return p


def build_english_cv() -> Document:
    doc = Document()
    configure_page(doc)

    # Header
    name = doc.add_paragraph()
    name.alignment = WD_ALIGN_PARAGRAPH.CENTER
    set_paragraph_format(name, before=0, after=2, line=1.0)
    r = name.add_run("Ziyi Sun")
    set_run_font(r, size=20, bold=True)

    subtitle = doc.add_paragraph()
    subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
    set_paragraph_format(subtitle, before=0, after=1, line=1.0)
    r = subtitle.add_run("B.Eng. Candidate · Ocean University of China")
    set_run_font(r, size=10.5, italic=True)

    contact = doc.add_paragraph()
    contact.alignment = WD_ALIGN_PARAGRAPH.CENTER
    set_paragraph_format(contact, before=0, after=2, line=1.05)
    r = contact.add_run(
        "sunziyi@stu.ouc.edu.cn  ·  sunziyiwiner@163.com  ·  "
        "https://ziyisun85-ops.github.io/"
    )
    set_run_font(r, size=9.5)

    objective = doc.add_paragraph()
    objective.alignment = WD_ALIGN_PARAGRAPH.CENTER
    set_paragraph_format(objective, before=0, after=2, line=1.05)
    r = objective.add_run(
        "Seeking Ph.D. opportunities in Embodied Intelligence and Robot Learning"
    )
    set_run_font(r, size=10, italic=True)

    # Education
    add_heading_line(doc, "EDUCATION")
    add_lr_line(doc, "Ocean University of China (985 / Double First-Class), Qingdao, China", "Sep. 2023 – Jun. 2027 (Expected)")
    add_body(
        doc,
        "B.Eng. in Mechanical Design, Manufacturing and Automation  |  GPA: 3.7/4.0  |  Rank: 1/56  |  Avg.: 89.2",
        bold=True,
        size=10,
    )
    add_body(
        doc,
        "Selected Courses: Advanced Mathematics (97), Linear Algebra (93.5), Probability & Mathematical Statistics (93), "
        "Complex Variables & Integral Transforms (94.5), Numerical Methods (96), Theory of Machines and Mechanisms (89), "
        "Mechanical Vibrations (95), Engineering Measurement (93), Mechanics of Materials (91), "
        "Modern Mechanical Design Theory and Methods (100), Underwater Robotics (97.2)",
        size=9.5,
        after=2,
    )

    # Internship
    add_heading_line(doc, "RESEARCH EXPERIENCE")
    add_lr_line(doc, "Roboparty Lab, Beijing, China", "Jul. 2026 – Oct. 2026")
    add_body(doc, "Frontier Research and Open Source Intern (Mentor: Jagger)", bold=True, size=10)
    add_bullet(
        doc,
        "Contribute to frontier research and open-source development for embodied intelligence and robot learning systems.",
    )

    # Research Interests
    add_heading_line(doc, "RESEARCH INTERESTS")
    add_body(
        doc,
        "My research goal is to build generalizable embodied agents that can perceive, reason, and act in complex "
        "real-world environments. I am particularly interested in:",
        size=10,
        after=1,
    )
    add_bullet(
        doc,
        "Robot Learning & Control: imitation/reinforcement learning, sim-to-real transfer, multimodal perception, dexterous manipulation",
    )
    add_bullet(
        doc,
        "Embodied AI: perception–decision–control loops, multimodal foundation models, world models, long-horizon task execution",
    )
    add_bullet(
        doc,
        "Human–Robot Interaction: learning from human demonstrations and enabling safe, adaptive robot behavior",
    )

    # Publications
    add_heading_line(doc, "PUBLICATIONS")
    add_body(
        doc,
        "[1] A Two-Stage Learning Framework for Autonomous Obstacle Avoidance of an Eel-Like Robot",
        bold=True,
        size=10,
        before=2,
    )
    add_body(
        doc,
        "IEEE International Conference on Real-time Computing and Robotics (IEEE RCAR), 2026  ·  Best Paper Finalist",
        italic=True,
        size=9.5,
    )
    add_body(
        doc,
        "Ziyi Sun, Luwen Li, Changhong He, Shuo Jiang, Zhuoyan Wang, Haoyuan Cheng*",
        size=9.5,
    )
    add_body(
        doc,
        "Proposed a two-stage learning framework combining multimodal perception, expert demonstration, PPO, "
        "AHTA-CPG, and domain randomization to address propulsion–steering coupling and sim-to-real gap for an "
        "eel-like robot; achieved 70%–90% obstacle-avoidance success over 60 real water-tank trials.",
        size=9.5,
        after=3,
    )

    add_body(
        doc,
        "[2] CHEROE: Every Humanoid Skill as a Traject",
        bold=True,
        size=10,
    )
    add_body(doc, "Under Review, AAAI 2027 (CCF-A)", italic=True, size=9.5)
    add_body(
        doc,
        "Ziyi Sun, Jingwen Chen, Yuxi Wang, Xiuze Xia, Long Cheng, Zhaoxiang Zhang, Junyu Dong",
        size=9.5,
    )
    add_body(
        doc,
        "Designed a unified SkillMotion representation and Tracker interface with boundary compatibility assessment, "
        "trajectory fusion, and bridging-action retrieval for reusable multi-source humanoid skills and continuous execution.",
        size=9.5,
        after=3,
    )

    add_body(
        doc,
        "[3] A Streaming Evaluation Framework for Deep-Sea Mining: Comparing Centralized and Distributed Nodule Collection",
        bold=True,
        size=10,
    )
    add_body(
        doc,
        "Under Review, Journal of Marine Science and Engineering (JMSE)",
        italic=True,
        size=9.5,
    )
    add_body(doc, "Ziyi Sun, Changhong He", size=9.5)
    add_body(
        doc,
        "Built a unified simulation and benchmarking pipeline for scalable robot learning, integrating CAD-based robot "
        "models and reproducible comparisons of centralized vs. distributed multi-AUV nodule collection.",
        size=9.5,
        after=3,
    )

    add_body(
        doc,
        "[4] Design of Biomimetic Robotic Eel with Omnidirectional Wire-Driven Body and Controlled Central Pattern Generator",
        bold=True,
        size=10,
    )
    add_body(
        doc,
        "IEEE International Conference on Computer, Control and Robotics (ICCCR), 2025",
        italic=True,
        size=9.5,
    )
    add_body(
        doc,
        "Presented the mechanical design and control framework of a biomimetic robotic eel with an omnidirectional "
        "quadruple wire-driven body and CPG-based locomotion control.",
        size=9.5,
        after=2,
    )

    # Projects
    add_heading_line(doc, "SELECTED PROJECTS")
    add_lr_line(
        doc,
        "Variable-Length Multimodal Biomimetic Robotic Fish for Deep-Sea Inspection (OUC SRDP)",
        "Dec. 2025 – Dec. 2026",
    )
    add_body(doc, "Project Leader", bold=True, size=10)
    add_body(
        doc,
        "Designed a biomimetic robotic fish with continuously adjustable body length (1.2–1.5 m) and fused school-fish "
        "visual tracking with water-quality sensing for joint monitoring and early warning in deep-sea cage aquaculture.",
        size=9.5,
        after=3,
    )

    add_lr_line(
        doc,
        "Cross-Domain Wire-Driven Biomimetic Eel Robot (Provincial Innovation Program)",
        "Dec. 2024 – Nov. 2025",
    )
    add_body(doc, "Project Leader", bold=True, size=10)
    add_body(
        doc,
        "Developed an eel-like robot with cruising and active depth-control capabilities for floating offshore wind "
        "dynamic-cable inspection, integrating visual tracking, Gaussian Splatting reconstruction, and damage detection.",
        size=9.5,
        after=3,
    )

    add_lr_line(
        doc,
        "High-Stability Duck-Type Wave Energy Converter with Pendulum PTO (National Innovation Program)",
        "Dec. 2024 – Nov. 2025",
    )
    add_body(doc, "Core Member", bold=True, size=10)
    add_body(
        doc,
        "Improved capture efficiency and output stability via AQWA hydrodynamics modeling, genetic-algorithm shape "
        "optimization, and a combined PTO with nonlinear adaptive regulation under complex sea states.",
        size=9.5,
        after=2,
    )

    # Patents
    add_heading_line(doc, "PATENTS & SOFTWARE COPYRIGHTS")
    add_bullet(
        doc,
        "Invention Patent: A High-Stability PTO System for Duck-Type Wave Energy Converters (First Student Inventor)",
    )
    add_bullet(
        doc,
        "Software Copyright: Underwater Polarized Light Field Prediction Software for Biomimetic Navigation V1.0 "
        "(Reg. No. 2024SR0140235)",
    )

    # Honors
    add_heading_line(doc, "HONORS & AWARDS")
    add_bullet(doc, "Shandong Provincial Government Scholarship (2025)")
    add_bullet(doc, "First-Class Comprehensive Scholarship, Ocean University of China (2024 & 2025)")
    add_bullet(doc, "Outstanding Student / Outstanding Student Cadre, Ocean University of China (2024 & 2025)")
    add_bullet(doc, "National First Prize, National 3D Digital Innovation Design Competition (2025 & 2026)")
    add_bullet(doc, "National Second Prize, Contemporary Undergraduate Mathematical Contest in Modeling (CUMCM, 2025)")
    add_bullet(doc, "Honorable Mention, Mathematical Contest in Modeling (MCM, 2025)")
    add_bullet(doc, "National Third Prize, 18th National College Student Energy Conservation Competition (2024)")
    add_bullet(doc, "National Gold Award, 3rd “Chuangyi Cup” National Undergraduate Extracurricular Academic Contest (2024)")
    add_bullet(doc, "National Second Prize, 14th Marine Vehicle Design and Production Competition (2025)")
    add_bullet(doc, "National First Prize, CRRC Cup National Undergraduate Renewable Energy Competition (2026)")
    add_bullet(doc, "Science and Technology Innovation Pioneer Team, College of Engineering, OUC (2024)")

    # Leadership
    add_heading_line(doc, "LEADERSHIP & SERVICE")
    add_bullet(doc, "Class Representative / Class Monitor, Ocean University of China")
    add_bullet(doc, "Team Leader, College of Engineering Debate Team")
    add_bullet(
        doc,
        "Director of Event Organization, Student Association for the Study of Socialism with Chinese Characteristics",
    )
    add_bullet(doc, "Outstanding Individual in Summer Social Practice, Ocean University of China")
    add_bullet(doc, "Volunteer, 2023 Qingdao Marathon; Member, 2024 National College Student Volunteer Outreach Team")

    # Skills
    add_heading_line(doc, "SKILLS")
    add_bullet(doc, "Robot Platform: Unitree G1")
    add_bullet(doc, "Programming & Tools: Python, C/C++, MATLAB, PyTorch, Isaac Gym/Lab, MuJoCo")
    add_bullet(doc, "Mechanical Design: SolidWorks, AutoCAD")
    add_bullet(doc, "Languages: Chinese (Native); English (CET-4: 591, CET-6: 482)")

    return doc


def build_chinese_cv() -> Document:
    doc = Document()
    configure_page(doc)
    EA = "SimSun"
    EA_H = "SimHei"

    name = doc.add_paragraph()
    name.alignment = WD_ALIGN_PARAGRAPH.CENTER
    set_paragraph_format(name, before=0, after=2, line=1.0)
    r = name.add_run("孙梓轶")
    set_run_font(r, size=20, bold=True, east_asia=EA_H)
    r2 = name.add_run("  /  Ziyi Sun")
    set_run_font(r2, size=14, bold=True, east_asia=EA_H)

    subtitle = doc.add_paragraph()
    subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
    set_paragraph_format(subtitle, before=0, after=1, line=1.0)
    r = subtitle.add_run("中国海洋大学 · 机械设计制造及其自动化 · 本科在读")
    set_run_font(r, size=10.5, italic=True, east_asia=EA)

    contact = doc.add_paragraph()
    contact.alignment = WD_ALIGN_PARAGRAPH.CENTER
    set_paragraph_format(contact, before=0, after=2, line=1.05)
    r = contact.add_run(
        "sunziyi@stu.ouc.edu.cn  ·  sunziyiwiner@163.com  ·  https://ziyisun85-ops.github.io/"
    )
    set_run_font(r, size=9.5, east_asia=EA)

    objective = doc.add_paragraph()
    objective.alignment = WD_ALIGN_PARAGRAPH.CENTER
    set_paragraph_format(objective, before=0, after=2, line=1.05)
    r = objective.add_run("申请方向：具身智能 / 机器人学习 相关博士项目")
    set_run_font(r, size=10, italic=True, east_asia=EA)

    add_heading_line(doc, "教育经历", east_asia=EA_H)
    add_lr_line(
        doc,
        "中国海洋大学（985 / 双一流 A 类）",
        "2023.09 – 2027.06（预计）",
        east_asia=EA,
    )
    add_body(
        doc,
        "机械设计制造及其自动化（ESI Top 1‰）  |  综合排名：1/56  |  GPA：3.7/4.0  |  均分：89.2",
        bold=True,
        size=10,
        east_asia=EA,
    )
    add_body(
        doc,
        "核心课程：高等数学(97)、线性代数(93.5)、概率论与数理统计(93)、复变函数与积分变换(94.5)、"
        "计算方法及其应用(96)、机械原理(89)、机械振动基础(95)、工程测试技术(93)、材料力学(91)、"
        "现代机械设计理论与方法(100)、水下机器人技术(97.2)",
        size=9.5,
        after=2,
        east_asia=EA,
    )

    add_heading_line(doc, "科研经历", east_asia=EA_H)
    add_lr_line(doc, "Roboparty Lab（北京）", "2026.07 – 2026.10", east_asia=EA)
    add_body(doc, "前沿研究与开源实习（导师：Jagger）", bold=True, size=10, east_asia=EA)
    add_bullet(
        doc,
        "参与具身智能与机器人学习方向的前沿研究及开源项目开发。",
        east_asia=EA,
    )

    add_heading_line(doc, "研究方向", east_asia=EA_H)
    add_body(
        doc,
        "致力于构建可泛化的具身智能体，使其在复杂真实环境中实现感知、决策与执行。重点关注：",
        size=10,
        east_asia=EA,
    )
    add_bullet(
        doc,
        "机器人学习与控制：模仿学习/强化学习、仿真到现实迁移、多模态感知、灵巧操作",
        east_asia=EA,
    )
    add_bullet(
        doc,
        "具身智能：感知—决策—控制闭环、多模态基础模型、世界模型、长时程任务执行",
        east_asia=EA,
    )
    add_bullet(
        doc,
        "人机交互：从人类示范中学习，实现安全、自适应的机器人行为",
        east_asia=EA,
    )

    add_heading_line(doc, "学术论文", east_asia=EA_H)
    add_body(
        doc,
        "[1] A Two-Stage Learning Framework for Autonomous Obstacle Avoidance of an Eel-Like Robot",
        bold=True,
        size=10,
        before=2,
        east_asia=EA,
    )
    add_body(
        doc,
        "IEEE RCAR 2026（CAA/B）· Best Paper Finalist",
        italic=True,
        size=9.5,
        east_asia=EA,
    )
    add_body(
        doc,
        "Ziyi Sun, Luwen Li, Changhong He, Shuo Jiang, Zhuoyan Wang, Haoyuan Cheng*",
        size=9.5,
        east_asia=EA,
    )
    add_body(
        doc,
        "针对仿鳗鱼机器人推进—转向耦合及仿真到真实的域偏移问题，融合多模态感知、专家示范、PPO、AHTA-CPG 与域随机化，"
        "构建自主避障框架；在 60 次实体水槽测试中实现 70%–90% 的避障成功率。",
        size=9.5,
        after=3,
        east_asia=EA,
    )

    add_body(
        doc,
        "[2] CHEROE: Every Humanoid Skill as a Traject",
        bold=True,
        size=10,
        east_asia=EA,
    )
    add_body(doc, "在投，AAAI 2027（CCF/A）", italic=True, size=9.5, east_asia=EA)
    add_body(
        doc,
        "Ziyi Sun, Jingwen Chen, Yuxi Wang, Xiuze Xia, Long Cheng, Zhaoxiang Zhang, Junyu Dong",
        size=9.5,
        east_asia=EA,
    )
    add_body(
        doc,
        "针对异构人形机器人技能难以复用与连续组合的问题，设计统一的 SkillMotion 表示与 Tracker 执行接口，"
        "并通过边界兼容性评估、轨迹融合和桥接动作检索，实现多源技能的统一管理、稳定执行与连续衔接。",
        size=9.5,
        after=3,
        east_asia=EA,
    )

    add_body(
        doc,
        "[3] A Streaming Evaluation Framework for Deep-Sea Mining: Comparing Centralized and Distributed Nodule Collection",
        bold=True,
        size=10,
        east_asia=EA,
    )
    add_body(
        doc,
        "在投，Journal of Marine Science and Engineering (JMSE)",
        italic=True,
        size=9.5,
        east_asia=EA,
    )
    add_body(doc, "Ziyi Sun, Changhong He", size=9.5, east_asia=EA)
    add_body(
        doc,
        "构建面向可扩展机器人学习的统一仿真与评测流水线，集成 CAD 机器人模型，并对集中式与分布式多 AUV "
        "结核采集方案进行可复现对比。",
        size=9.5,
        after=3,
        east_asia=EA,
    )

    add_body(
        doc,
        "[4] Design of Biomimetic Robotic Eel with Omnidirectional Wire-Driven Body and Controlled Central Pattern Generator",
        bold=True,
        size=10,
        east_asia=EA,
    )
    add_body(doc, "ICCCR 2025", italic=True, size=9.5, east_asia=EA)
    add_body(
        doc,
        "提出仿生机器鳗的机械设计与控制框架，融合全向四线驱动本体与可控中枢模式发生器，实现柔性水下运动。",
        size=9.5,
        after=2,
        east_asia=EA,
    )

    add_heading_line(doc, "项目经历", east_asia=EA_H)
    add_lr_line(
        doc,
        "面向深海探测的可变长多模态感知仿生机器鱼（中国海洋大学 SRDP）",
        "2025.12 – 2026.12",
        east_asia=EA,
    )
    add_body(doc, "项目负责人", bold=True, size=10, east_asia=EA)
    add_body(
        doc,
        "面向深海网箱巡航与狭窄空间巡检需求，设计体长可在 1.2–1.5 m 无级调节的仿生机器鱼，并融合鱼群视觉跟踪与水质传感信息，"
        "实现对鱼群行为及养殖环境风险的联合监测与预警。",
        size=9.5,
        after=3,
        east_asia=EA,
    )

    add_lr_line(
        doc,
        "基于拉线式结构的可跨域仿生鳗鲡机器人（省级大学生创新创业项目）",
        "2024.12 – 2025.11",
        east_asia=EA,
    )
    add_body(doc, "项目负责人", bold=True, size=10, east_asia=EA)
    add_body(
        doc,
        "面向深远海浮式风电动态缆巡检难题，设计具备巡航与主动升潜能力的仿鳗鲡机器人，并融合视觉跟踪、"
        "Gaussian Splatting 三维重建与损伤检测，实现动态缆建模、缺陷定位及风险预警。",
        size=9.5,
        after=3,
        east_asia=EA,
    )

    add_lr_line(
        doc,
        "基于摆式 PTO 的高稳定鸭式波浪能发电装置（国家级大学生创新创业项目）",
        "2024.12 – 2025.11",
        east_asia=EA,
    )
    add_body(doc, "核心成员", bold=True, size=10, east_asia=EA)
    add_body(
        doc,
        "针对鸭式波浪能装置捕能效率与输出稳定性不足的问题，结合 AQWA 水动力建模、遗传算法外形优化、组合式 PTO 结构及"
        "非线性自适应调节，拓宽有效谐振频带并提升复杂波况下的能量转换效率。",
        size=9.5,
        after=2,
        east_asia=EA,
    )

    add_heading_line(doc, "专利与软著", east_asia=EA_H)
    add_bullet(
        doc,
        "发明专利：一种高稳定鸭式波浪能转换装置的 PTO 系统（学生第一发明人）",
        east_asia=EA,
    )
    add_bullet(
        doc,
        "软件著作权：面向仿生导航的水下偏振光场预测软件 V1.0（登记号：2024SR0140235）",
        east_asia=EA,
    )

    add_heading_line(doc, "主要荣誉", east_asia=EA_H)
    add_bullet(doc, "山东省政府奖学金（2025）", east_asia=EA)
    add_bullet(doc, "中国海洋大学综合类奖学金一等奖（2024、2025）", east_asia=EA)
    add_bullet(doc, "中国海洋大学优秀学生 / 优秀学生干部（2024、2025）", east_asia=EA)
    add_bullet(doc, "全国三维数字化创新设计大赛 国家一等奖（2025、2026）", east_asia=EA)
    add_bullet(doc, "全国大学生数学建模竞赛 国家二等奖（2025）", east_asia=EA)
    add_bullet(doc, "美国大学生数学建模竞赛 Honorable Mention（2025）", east_asia=EA)
    add_bullet(doc, "全国大学生节能减排社会实践与科技竞赛 国家三等奖（2024）", east_asia=EA)
    add_bullet(doc, "“创祎杯”全国大学生课外学术科技作品大赛 国家级金奖（2024）", east_asia=EA)
    add_bullet(doc, "海洋航行器设计与制作大赛 国家二等奖（2025）", east_asia=EA)
    add_bullet(doc, "“中国中车杯”全国大学生可再生能源优秀科技作品竞赛 国家一等奖（2026）", east_asia=EA)
    add_bullet(doc, "工程学院科技创新先锋团队（2024）", east_asia=EA)

    add_heading_line(doc, "学生工作与社会实践", east_asia=EA_H)
    add_bullet(doc, "班级代表 / 班长，中国海洋大学", east_asia=EA)
    add_bullet(doc, "工程学院辩论队队长", east_asia=EA)
    add_bullet(doc, "中国特色社会主义理论学习研究会活动组织部部长", east_asia=EA)
    add_bullet(doc, "中国海洋大学暑期社会实践优秀个人", east_asia=EA)
    add_bullet(doc, "2023 青岛马拉松志愿者；2024 全国大学生宣讲遵义会议精神志愿宣讲团成员", east_asia=EA)

    add_heading_line(doc, "专业技能", east_asia=EA_H)
    add_bullet(doc, "机器人平台：Unitree G1", east_asia=EA)
    add_bullet(doc, "编程与工具：Python、C/C++、MATLAB、PyTorch、Isaac Gym/Lab、MuJoCo", east_asia=EA)
    add_bullet(doc, "机械设计：SolidWorks、AutoCAD", east_asia=EA)
    add_bullet(doc, "语言能力：英语 CET-4：591；CET-6：482", east_asia=EA)

    return doc


def main():
    OUT_DIR.mkdir(parents=True, exist_ok=True)
    en_path = OUT_DIR / f"Ziyi_Sun_CV_EN_{DATE_TAG}.docx"
    zh_path = OUT_DIR / f"孙梓轶_简历_中文_{DATE_TAG}.docx"
    build_english_cv().save(en_path)
    build_chinese_cv().save(zh_path)
    print(f"Wrote {en_path}")
    print(f"Wrote {zh_path}")


if __name__ == "__main__":
    main()
